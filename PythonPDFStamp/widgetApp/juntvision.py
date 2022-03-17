import os
import re
import time
import pdfplumber
import pymysql
import xlwings as xw

from pathlib import Path
from docx import Document
from win32com import client as wc

class Juntvision():
    #把每一行的采购信息，汇总添加到这个list中
    buy_order_list = []

    #获取海康采购订单的详细信息
    def get_buy_order_info(self, buy_order_path):
        pdf = pdfplumber.open(buy_order_path)
        pdf_page_one = pdf.pages[0]
        page_one_text = pdf_page_one.extract_text()
        page_text_list = page_one_text.split()
        #获取海康采购合同订购单号
        buy_order_id = page_text_list[8].split('：')[1]

        #获取海康采购合同日期
        buy_date = page_text_list[9].split('：')[1]

        #获取海康公司名称
        supplier_name = page_text_list[14].split('：')[1]

        #截取订购明细信息表
        page_one_table = pdf_page_one.extract_table()[1:-1]
        #print(page_one_table)

        for product in page_one_table:
            #用于储存采购的每一行信息，然后插入数据库中
            product_info_dic = {}
            product_info_dic['合同编号'] = buy_order_id
            product_info_dic['签订日期'] = buy_date
            product_info_dic['供货方'] = supplier_name
            #print(product)
            product_name = product[1]
            product_info_dic['商品名称'] = product_name
            #print(product_name)
            product_model = product[2]
            if '\n' in product_model:
                product_model = product_model.replace('\n', '')
            product_info_dic['型号'] = product_model
            #print(product_model)
            quantity = product[3]
            product_info_dic['数量'] = quantity
            #print(quantity)
            unit_price = product[4]
            product_info_dic['单价'] = unit_price
            #print(unit_price)
            self.buy_order_list.append(product_info_dic)
            #print(product_info_dic)
        print(len(self.buy_order_list))
        print(self.buy_order_list)
        return self.buy_order_list
    #获取中科院自动化所销售订单明细
    def get_sale_order_info_CASIA(self,path):
        try:
            sale_order_info_all = []
            doc = Document(path)
            filepath,filename =os.path.split(path)
            #获取合同编号
            print(filename)
            #order_number_str = filter(str.isdigit,filename)
            order_number_str =re.findall(r'\d+',filename)[0]
            #获取订购日期
            date_str = order_number_str[0:8]

            tables = doc.tables
            table_two = tables[1]

            for row in table_two.rows[1:]:
                    sale_order_info_dic = {}
                    if row.cells[4].text.isdigit():
                        sale_order_info_dic['购货单位'] = '中国科学院自动化研究所'
                        sale_order_info_dic['合同编号'] = order_number_str
                        sale_order_info_dic['签订日期'] = date_str
                        if '年'in sale_order_info_dic['签订日期']:
                            date_list =re.split('年|月|日',sale_order_info_dic['签订日期'])
                            sale_order_info_dic['签订日期'] ='-'.join(date_list)
                        if '_' in sale_order_info_dic['签订日期']:
                            sale_order_info_dic['签订日期']=sale_order_info_dic['签订日期'].replace('_','')
                        sale_order_info_dic['商品名称'] = row.cells[1].text
                        sale_order_info_dic['型号'] = row.cells[2].text
                        sale_order_info_dic['数量'] = row.cells[5].text
                        sale_order_info_dic['单价'] = row.cells[4].text
                        sale_order_info_all.append(sale_order_info_dic)
                    else:
                        continue
        except Exception as ex:
                print('提取销售订单信息错误{}'.format(path))
                print(ex.args)
        return sale_order_info_all
    
    #获取君泰销售订单明细word模板的合同
    def get_sale_order_info_word(self, path):
        '''
        :param path 传入的docx的销售订单文件路径
        :return  sale_order_info_all 一个list 装了需要导入到数据库的销售明细信息
        '''
        try:
            doc = Document(path)
            tables = doc.tables
            table_one = tables[0]
            #print(type(table_one._cells))
            sale_order_info_list = []
            sale_order_info_all = []
            for cell in table_one._cells:
                # print(type(table_one._cells))
                #sale_order_info_dic={}
                sale_order_info_list.append(cell.text.split('：')[1])
                #print(cell.text)
            #获取购货单位和合同编号
            sale_order_info_list = sale_order_info_list[0:2]
            print(sale_order_info_list)
            table_three = tables[2]
            #获取合同签订日期
            print(table_three._cells[-1].text)
            sale_order_info_list.append(table_three._cells[-1].text.split('：')[1])
            print(sale_order_info_list)
            table_two = tables[1]

            for row in table_two.rows[1:-1]:
                sale_order_info_dic = {}
                if row.cells[4].text.isdigit():
                    sale_order_info_dic['购货单位'] = sale_order_info_list[0]
                    sale_order_info_dic['合同编号'] = sale_order_info_list[1]
                    sale_order_info_dic['签订日期'] = table_three._cells[-1].text.split(
                        '：')[1]
                    if '年'in sale_order_info_dic['签订日期']:
                        date_list =re.split('年|月|日',sale_order_info_dic['签订日期'])
                        sale_order_info_dic['签订日期'] ='-'.join(date_list)
                    if '_' in sale_order_info_dic['签订日期']:
                        sale_order_info_dic['签订日期']=sale_order_info_dic['签订日期'].replace('_','')
                    sale_order_info_dic['商品名称'] = row.cells[1].text
                    sale_order_info_dic['型号'] = row.cells[2].text
                    sale_order_info_dic['数量'] = row.cells[4].text
                    sale_order_info_dic['单价'] = row.cells[5].text
                    sale_order_info_all.append(sale_order_info_dic)
                else:
                    continue
        except :
            print('提取销售订单信息错误{}'.format(path))
            
        return sale_order_info_all

    #获取君泰销售订单明细Excel模板的合同
    def get_sale_order_info_excel(self,path):
        sale_order_info_all =[]
        app = xw.App(visible=False, add_book=False)
        app.display_alerts = False
        app.screen_updating = False
        wb = app.books.open(path)
        sheet = wb.sheets['Sheet1']
        # 获取有效的excel的区域
        info = sheet.used_range
        # 获取最后一行的行数
        nrows = info.rows.count
        print(nrows)
        filepath,filename =os.path.split(path)
      
        print(filename)
        #order_number_str = filter(str.isdigit,filename)
        order_number_str =re.findall(r'\d+',filename)[0]
        #获取订购日期
        date_str = order_number_str[0:8]

        # 获取货物名称，规格型号，单价，数量
        row_num =11
        #从第11行开始，判断非空的产品行有几行
        while sheet.range("G" +str(row_num)).value != None:
            sale_order_info_dic = {}
            sale_order_info_dic['购货单位'] ="".join(sheet.range('D4').value.split()) 
            sale_order_info_dic['合同编号'] = sheet.range('H4').value
            sale_order_info_dic['签订日期'] = date_str  
            sale_order_info_dic['商品名称'] = sheet.range("C"+str(row_num)).value
            sale_order_info_dic['型号'] = sheet.range("D"+str(row_num)).value
            sale_order_info_dic['数量'] = sheet.range("F"+str(row_num)).value
            sale_order_info_dic['单价'] = sheet.range("G"+str(row_num)).value
            sale_order_info_all.append(sale_order_info_dic)
            row_num+=1
        print(sale_order_info_all)
        wb.save()
        wb.close()
        app.quit()
        return sale_order_info_all

    #把海康采购的订单数据录入到数据库中
    def insert_buy_order(self, buy_order_list):
        connection = pymysql.connect(host='localhost',
                                     user='root',
                                     passwd='11251125',
                                     database='juntevision')
        cursor = connection.cursor()
        for product_item in buy_order_list:
            insert_sql="insert into buydetail(buyOrderID, buyDate, supplierName, productName, \
                productModel, quantity, unitPrice) values('%s','%s','%s','%s','%s','%d','%f')"                                                                                               % \
                    (product_item['合同编号'],product_item['签订日期'],product_item['供货方'],product_item['商品名称'], \
                        product_item['型号'],int(product_item['数量']),float(product_item['单价']))
            try:
                cursor.execute(insert_sql)
                connection.commit()
            except:
                connection.rollback()
                print("插入数据异常{}".format(buy_order_list))
               
        connection.close()

    def insert_saledetail_table(self, sale_order_list):
        connection = pymysql.connect(host='localhost',
                                     user='root',
                                     passwd='11251125',
                                     database='juntevision')
        cursor = connection.cursor()
        for product_item in sale_order_list:
            #在做数据添加的时候，要把已经添加了的合同号做删除
            select_sql = "select saleOrderID from saledetail where saleOrderID = '%s'" % product_item[
                '合同编号']
            result = cursor.execute(select_sql)
            print(result)
            if result < len(sale_order_list):
                print(product_item)
                try:
                    insert_sql="insert into saledetail(saleOrderID, customerName, saleDate, productName, productModel, \
                        quantity,unitPrice) values('%s','%s','%s','%s','%s','%d','%f')" % \
                            (product_item['合同编号'],product_item['购货单位'],product_item['签订日期'],product_item['商品名称'], \
                                product_item['型号'],int(product_item['数量']),float(product_item['单价']))
                except ValueError as e:
                    print(e.args)
                    pass
                    continue
                try:
                    cursor.execute(insert_sql)
                    connection.commit()
                    print('客户{}合同编号{}成功插入销售订单数据到数据库中'.format(product_item['购货单位'],product_item['合同编号']))
                except pymysql.Error as e:
                    connection.rollback()
                    print("\033[31m插入数据异常{}\033[0m".format(sale_order_list))
                    print(e.args)
                    print("\033[31m有问题的销售合同信息list：{}\033[0m".format(sale_order_list))
            elif result == len(sale_order_list):
                print("销售合同已存在数据库中...")
        connection.close()

    #把海康合同中的产品信息录入到数据库中
    #buy_order_list是从海康订购合同中提取的商品信息
    def insert_product_table(self, buy_order_list):
        connection = pymysql.connect(host='localhost',
                                     user='root',
                                     passwd='11251125',
                                     database='juntevision')
        cursor = connection.cursor()
        for product_item in buy_order_list:
            select_sql = "select productModel from product where productModel = '%s'" % product_item[
                '型号']
            result = cursor.execute(select_sql)
            print(result)
            if result == 0:
                insert_sql="insert into product(productName, productModel, productBuyPrice, supplierName \
                    ) values('%s','%s','%f','%s')"                                                   % \
                        (product_item['商品名称'],product_item['型号'],float(product_item['单价']),product_item['供货方'])
                try:
                    cursor.execute(insert_sql)
                    connection.commit()
                except:
                    connection.rollback()
                    print("\033[31m插入数据异常\033[0m")
        connection.close()

    #修改添加完后的采购订单文件名，后面加上mysql字符
    def rename_buy_order(self, buy_order_path):
        #file_name = os.path.basename(buy_order_path)
        path = Path(buy_order_path)
        file_stem = path.stem
        file_new_stem = file_stem + '-mysql'
        file_new_name = str(path.parent / file_new_stem) + path.suffix
        print(file_new_name)
        path.rename(file_new_name)

    #把时间戳转化为时间: 1479264792 to 2016-11-16 10:53:12
    # def TimeStampToTime(self,timestamp):
	#     timeStruct = time.localtime(timestamp)
	#     return time.strftime('%Y-%m-%d %H:%M:%S',timeStruct)
    
    # #'''获取文件的修改时间'''
    # def get_FileModifyTime(self,filePath):
	#     t = os.path.getmtime(filePath)
	#     return self.TimeStampToTime(t)


    #君泰通达销售订单目录筛选
    #20220310-做到把销售合同目录下的客户合同目录按照最新修改时间排序
    #下一步遍历文件夹内的文件，提取合同内的信息添加到数据库中
    def get_sale_folder_path(self,folder_path):
        '''
        :param folder_path:传入文件夹路径
        :return folder_list:返回2022年1月1日之后的修改文件夹的list
        '''
        folder_list=[]
        folder_path_names = os.listdir(folder_path) #返回的是folder_path里面的所有目录的路径list，里面的路径是str
        folder_path_names.sort(key=lambda folder:os.path.getmtime(os.path.join(folder_path,folder)),reverse=True)
        for folder_path_name in folder_path_names:
            path = os.path.join(folder_path,folder_path_name)            
            #print(os.path.getctime(file_path))#获取创建时间，返回值是标准时间，需要转换
            time_secs = os.path.getmtime(path)#获取最后修改时间，返回值是标准时间，需要转换
            time_local = time.localtime(time_secs) #返回的是struct_time类型，用struct_time类型来比较时间
            time_2021_start ='01/01/2022'
            time_start = time.strptime(time_2021_start,"%d/%m/%Y")
            #选区修改时间，也就是新增文件时间大于2020年1月1日的目录
            if time_local>time_start:
                folder_list.append(path)
                #print(path)
            #print(time_start)
            #time_str =time.strftime('%Y-%m-%d %H:%M:%S',time_local)
            #print(time_str)
            #print(os.path.getatime(file_path))#获取访问时间，返回值是标准时间，需要转换
            
        return folder_list
    
    #找到文件夹下所有doc文件的绝对路径，并返回路径列表
    def get_sale_file_path_with_filter(self,foler_path,filter=['.doc']):
        '''
        :param foler_path:传入的文件夹路径
        :param filter:传入的过滤list,用于筛选特定后缀文件,比如：.doc,.docx,
        :return file_path_list :默认传出doc文件路径列表
        '''
        file_path_list =[]
        filter =filter
        for root,dirs,files in os.walk(foler_path):
            for file in files:
                file_path = os.path.join(root,file)
                ext = os.path.splitext(file_path)[1] #获取文件后缀
                #判断文件后缀，判断文件名称是否包含“购销合同”这几个字符
                if ext in filter and '购销合同' in file_path:
                    file_path_list.append(file_path)
        #print(file_path_list)
        return file_path_list
    

    def convert_doc_to_docx(self,file_path_list):
        #print(file_path_list) 
        #将文件名和后缀分割   
        for doc_path in file_path_list:
            split_name =os.path.splitext(doc_path)
            docx_file_path =split_name[0]+'.docx'
        #print(docx_file_path)     
        #如果docx文件不存在，则另存doc文件到docx文件
            if doc_path.endswith('.doc') and not doc_path.startswith('~$'):
                #print(doc_path)
                if os.path.exists(docx_file_path):
                    continue
                else:
                    word =wc.Dispatch("kwps.Application")
                    #word.Visible =0
                    #word.DisplayAlerts =0          
                    doc = word.Documents.Open(doc_path)                                 
                    doc.SaveAs(split_name[0]+'.docx',12) #12表示docx格式
                    time.sleep(2)
                    doc.Close()
                    #word.Visible =1
                    #word.DisplayAlerts =1  
                    word.Quit()
        
if __name__ == "__main__":
    #test_path = 'H:\\崔向阳坚果云\\销售合同\\淄博蓝达智能视觉科技有限公司\\1月\\202201131009-淄博蓝达智能视觉科技有限公司采购合同.xlsx'
    juntObj = Juntvision()
    sale_folder_path ='H:\\崔向阳坚果云\\销售合同'   
    folder_path_list =juntObj.get_sale_folder_path(sale_folder_path)
    for folder_path in folder_path_list:
         file_path_list = juntObj.get_sale_file_path_with_filter(folder_path,['.xlsx'])
         print(file_path_list)
         if len(file_path_list)>0:
             for file_path in file_path_list:
                 sale_info =juntObj.get_sale_order_info_excel(file_path)
                 juntObj.insert_saledetail_table(sale_info)

            
        #print(len(file_path_list))
        #if len(file_path_list) !=0:
        #print(file_path_list)
            #juntObj.convert_doc_to_docx(file_path_list)
    
    #print(folder_path_list[0])

    #sale_order_list = juntObj.get_sale_order_info(path)
    #juntObj.insert_saledetail_table(sale_order_list)

    # folder_path ='H:\\崔向阳坚果云\\进货合同\\0杭州海康智能科技有限公司\\1月\\'
    # file_path_list =[]
    # for file in os.listdir(folder_path):
    #     file_path = os.path.join(folder_path,file)
    #     file_path_list.append(file_path)
    # print(file_path_list)
    # pdf_path ='C:\\Users\\郑勋\\Desktop\\2022669301北京君泰通达科技有限公司购销合同1.5.pdf'
    # juntObj =Juntvision()
    # #juntObj.rename_buy_order(pdf_path)
    # buy_order_list=juntObj.get_buy_order_info(pdf_path)
    # #juntObj.insert_buy_order(buy_order_list)
    # juntObj.insert_buy_order_product(buy_order_list)